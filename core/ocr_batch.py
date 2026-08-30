"""ocr_batch — OCR 批量识别（图片 + PDF）。

图片直接走 core.ocr_tool.ocr_image（RapidOCR）；PDF 用 PyMuPDF
逐页渲染为临时 PNG 后逐页识别，文本合并写回同名 .txt。
识别结果为空视为失败；不依赖 GUI。
"""
import os
import tempfile

from core.ocr_tool import ocr_image
from core.ocr_table import _box_point_size

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff", ".tif"}
PDF_EXTS = {".pdf"}
OCR_EXTS = IMAGE_EXTS | PDF_EXTS


def _report_error(progress_cb, message):
    if progress_cb:
        progress_cb(-1, message)


def _validate_conversion(input_path, output_path, output_ext, progress_cb):
    """在启动耗时识别前拒绝无效输入，避免误覆盖源文件。"""
    if not os.path.isfile(input_path):
        _report_error(progress_cb, "错误: 找不到输入文件")
        return False
    if not is_supported(input_path):
        _report_error(progress_cb, "错误: 不支持的输入文件格式")
        return False
    if os.path.splitext(output_path or "")[1].lower() != output_ext:
        _report_error(progress_cb, f"错误: 输出文件必须为 {output_ext} 格式")
        return False
    source = os.path.normcase(os.path.abspath(input_path))
    target = os.path.normcase(os.path.abspath(output_path))
    if source == target:
        _report_error(progress_cb, "错误: 输出文件不能覆盖输入文件")
        return False
    return True


def _staged_output_path(output_path, suffix):
    """在目标目录创建临时名，保证最终 os.replace 不跨文件系统。"""
    output_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(output_dir, exist_ok=True)
    fd, staged_path = tempfile.mkstemp(
        prefix=".fm_ocr_", suffix=suffix, dir=output_dir)
    os.close(fd)
    os.remove(staged_path)
    return staged_path


def _scaled_progress(progress_cb):
    """为最终校验与原子提交预留进度，同时原样传递错误和取消。"""
    if progress_cb is None:
        return None

    def report(value, message):
        progress_cb(value if value < 0 else min(90, int(value * 0.9)), message)

    return report


def is_supported(path):
    """是否为支持的 OCR 输入（图片或 PDF）。"""
    return os.path.splitext(path or "")[1].lower() in OCR_EXTS


def _ocr_pdf(input_path, lang, progress_cb=None):
    """逐页 OCR PDF，返回合并文本（每页之间空行分隔）。"""
    try:
        import pymupdf
    except ImportError:
        if progress_cb:
            progress_cb(-1, "错误: PyMuPDF (fitz) 未安装，无法识别 PDF")
        return ""
    try:
        doc = pymupdf.open(input_path)
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: PDF 打开失败 - {e}")
        return ""
    total = max(doc.page_count, 1)
    pages_text = []
    tmpdir = tempfile.mkdtemp(prefix="fm_ocr_")
    try:
        for idx in range(doc.page_count):
            if progress_cb:
                pct = int((idx + 1) / total * 95)
                progress_cb(pct, f"识别 PDF 第 {idx + 1}/{total} 页…")
            try:
                page = doc.load_page(idx)
                pix = page.get_pixmap(dpi=200)
                img_path = os.path.join(tmpdir, f"p{idx}.png")
                pix.save(img_path)
            except Exception:  # noqa: BLE001 - 单页渲染失败跳过
                continue
            text = ocr_image(img_path, lang)
            if text and text.strip():
                pages_text.append(text.strip())
    finally:
        try:
            doc.close()
        except Exception:  # noqa: BLE001
            pass
        try:
            for name in os.listdir(tmpdir):
                os.remove(os.path.join(tmpdir, name))
            os.rmdir(tmpdir)
        except OSError:
            pass
    if progress_cb:
        progress_cb(100, "识别完成")
    return "\n\n".join(pages_text)


def _sort_key(box):
    """按 行(y) → 列(x) 排序，用于按版面坐标重建文档。"""
    return (round(box[1] * 100), box[0])


def _cleanup_tmp(tmpdir):
    try:
        for name in os.listdir(tmpdir):
            os.remove(os.path.join(tmpdir, name))
        os.rmdir(tmpdir)
    except OSError:
        pass


def _denormalize_box(box, img_w, img_h):
    """把单个归一化 [0,1] 文本框反归一化为像素坐标 (x0,y0,x1,y1)。"""
    try:
        return (box[0] * img_w, box[1] * img_h, box[2] * img_w, box[3] * img_h)
    except (TypeError, ValueError, IndexError):
        return (0.0, 0.0, 0.0, 0.0)


def _denormalize_boxes(details, img_w, img_h):
    """把 ocr_image_detailed 的归一化 [0,1] 坐标反归一化为像素坐标。

    返回 [(text, (x0,y0,x1,y1)), ...]（图像像素坐标）。
    """
    boxes = []
    for t, b in details:
        if not t or not t.strip():
            continue
        box = _denormalize_box(b, img_w, img_h)
        if not box:
            continue
        boxes.append((t, box))
    return boxes


def _raw_to_details(raw, img_w, img_h):
    """把 rapidocr 原始结果（含 conf）转 [(text, 归一化box), ...]，供几何重建使用。"""
    details = []
    for item in raw or []:
        try:
            bbox, text = item[0], item[1]
            xs = [p[0] for p in bbox]
            ys = [p[1] for p in bbox]
            box = (min(xs) / max(img_w, 1), min(ys) / max(img_h, 1),
                   max(xs) / max(img_w, 1), max(ys) / max(img_h, 1))
            details.append((text, box))
        except Exception:  # noqa: BLE001 - 单条异常跳过
            continue
    return details


def _add_word_table(doc, table, grid_lines=None):
    """把一个重建表格写入 docx 为带边框的真·Word 表格（单元格可编辑）。

    - 单元格字号按原扫描件文本框高度换算得到的磅值设置；
    - 合并单元格：table['merges'] 记录跨行/跨列区域，生成 Word 合并；
    - 列宽：按 table['col_bounds'] 像素比例还原列宽；
    - 边框：默认带框（Table Grid）；仅当栅格线检测明确无框线时去框。
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    grid = table.get("grid", [])
    sizes = table.get("sizes", []) or []
    merges = table.get("merges", []) or []
    col_bounds = table.get("col_bounds", []) or []
    n_rows = len(grid)
    n_cols = max((len(r) for r in grid), default=0)
    if n_rows < 1 or n_cols < 1:
        return
    try:
        tbl = doc.add_table(n_rows, n_cols, style="Table Grid")
    except Exception:  # noqa: BLE001 - 个别 python-docx 版本无 Table Grid 样式
        tbl = doc.add_table(n_rows, n_cols)

    # 边框近似：默认带框；仅当 bbox 内明确无栅格线时才去框
    if not _table_has_border(table.get("bbox"), grid_lines):
        _remove_table_borders(tbl)

    # 先填充所有单元格文字与字号
    for r in range(n_rows):
        row_cells = tbl.rows[r].cells
        row_sizes = sizes[r] if r < len(sizes) else []
        for c in range(n_cols):
            val = grid[r][c] if c < len(grid[r]) else ""
            cell = row_cells[c]
            cell.text = ""  # 清空默认段落，便于设置字号
            para = cell.paragraphs[0]
            run = para.add_run(val or "")
            if c < len(row_sizes) and row_sizes[c] > 0:
                run.font.size = Pt(row_sizes[c])

    # 合并前，把合并区域内其它单元格的文字汇总到左上单元格，避免被吞
    for (r1, c1, r2, c2) in merges:
        if r2 >= n_rows or c2 >= n_cols:
            continue
        extra = []
        for rr in range(r1, r2 + 1):
            for cc in range(c1, c2 + 1):
                if rr == r1 and cc == c1:
                    continue
                t = grid[rr][cc] if cc < len(grid[rr]) else ""
                if t and t.strip():
                    extra.append(t)
        if extra:
            cur = grid[r1][c1]
            merged = (cur + " " + " ".join(extra)).strip() if cur else " ".join(extra)
            grid[r1][c1] = merged
            try:
                cell = tbl.cell(r1, c1)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = merged
            except Exception:  # noqa: BLE001
                pass

    # 合并单元格（面积降序，逐个尝试，冲突跳过）
    for (r1, c1, r2, c2) in sorted(
            merges, key=lambda m: -((m[2] - m[0] + 1) * (m[3] - m[1] + 1))):
        try:
            if r2 >= n_rows or c2 >= n_cols:
                continue
            tbl.cell(r1, c1).merge(tbl.cell(r2, c2))
        except Exception:  # noqa: BLE001
            pass

    # 列宽还原：按 bbox 各列像素比例缩放到表格总宽
    _apply_col_widths(tbl, col_bounds)

    # 表后留白
    doc.add_paragraph("")


def _table_has_border(bbox, grid_lines):
    """默认带框；仅当 bbox 内明确检测不到任何栅格线时才认为无框。"""
    if not grid_lines or not bbox:
        return True
    rows_y, cols_x = grid_lines
    if not rows_y or not cols_x:
        return True
    x0, y0, x1, y1 = bbox
    h = sum(1 for y in rows_y if y0 - 8 <= y <= y1 + 8)
    v = sum(1 for x in cols_x if x0 - 8 <= x <= x1 + 8)
    # 完全无横/纵线才视为无框，避免误删有框表格
    return not (h == 0 and v == 0)


def _remove_table_borders(tbl):
    """移除 Word 表格四周边框与内部网格线（用于无框表格）。"""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tblPr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "none")
            e.set(qn("w:sz"), "0")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
            borders.append(e)
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(borders)
    except Exception:  # noqa: BLE001
        pass


def _apply_col_widths(tbl, col_bounds):
    """按各列像素边界比例还原列宽，并固定表格总宽与布局。"""
    if not col_bounds or len(col_bounds) != len(tbl.columns):
        return
    widths = [max(1.0, x1 - x0) for (x0, x1) in col_bounds]
    total = sum(widths)
    if total <= 0:
        return
    table_width_pt = 6.3 * 72.0  # 与嵌入原图宽度一致
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        tblPr = tbl._tbl.tblPr
        w = tblPr.find(qn("w:tblW"))
        if w is None:
            w = OxmlElement("w:tblW")
            tblPr.append(w)
        w.set(qn("w:w"), str(int(table_width_pt)))
        w.set(qn("w:type"), "dxa")
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")
        for i, wd in enumerate(widths):
            try:
                tbl.columns[i].width = Pt(wd / total * table_width_pt)
            except Exception:  # noqa: BLE001
                pass
    except Exception:  # noqa: BLE001
        pass


def _build_docx_from_ocr(doc, tables, free_boxes, image_to_embed=None,
                         img_w=0, img_h=0, embed_width_inch=6.3, dpi=300,
                         image_path=None):
    """把「重建表格 + 自由文本」组装进 docx，保持阅读顺序。

    - image_to_embed 非 None 时先嵌入原图（保真，所见即所得）。
    - 之后按 y 坐标把表格与自由文本穿插排入（表格为可编辑真表格）。
    - dpi 用于把自由文本的字号还原为接近原文件的磅值。
    - image_path 用于栅格线检测，以近似还原边框（有框/无框）。
    """
    from docx.shared import Inches, Pt
    from core.ocr_table import lines_from_boxes, detect_grid_lines

    # 预检测整页栅格线，供各表格判断边框（仅一次读图）
    grid_lines = None
    if image_path:
        try:
            grid_lines = detect_grid_lines(image_path)
        except Exception:  # noqa: BLE001
            grid_lines = None

    if image_to_embed:
        try:
            doc.add_picture(image_to_embed, width=Inches(embed_width_inch))
        except Exception:  # noqa: BLE001 - 嵌图失败不影响文字
            pass

    blocks = []  # (y0, kind, payload)
    for t in tables:
        blocks.append((t.get("bbox", (0, 0, 0, 0))[1], "table", t))
    for line, size in lines_from_boxes(free_boxes, dpi=dpi):
        blocks.append((10 ** 9, "text", (line, size)))  # 自由文本统一排在表格之后兜底

    # 表格按 y 排序，自由文本插入：简单策略——先按表格出现顺序，
    # 自由文本聚到末尾（避免与表格 y 错位穿插导致混乱）。
    tables_sorted = sorted(
        [b for b in blocks if b[1] == "table"], key=lambda b: b[0])
    for _, _, t in tables_sorted:
        _add_word_table(doc, t, grid_lines)
    for _, _, payload in (b for b in blocks if b[1] == "text"):
        line, size = payload
        para = doc.add_paragraph()
        run = para.add_run(line)
        if size and size > 0:
            run.font.size = Pt(size)


def extract_docx_text(path):
    """提取 docx 全部文字（含表格单元格），用于结果回填预览。"""
    try:
        from docx import Document
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception:  # noqa: BLE001
        return ""


def _pdf_to_docx_via_pdf2docx(input_path, output_path, progress_cb=None):
    """数字 PDF 走 pdf2docx，保留表格/排版（仅转换已存在的文本层）。

    扫描型页面在 pdf2docx 下会变成图片（不可编辑），故仅用于数字文档。
    """
    import sys
    import pymupdf
    # pdf2docx 内部 import fitz，预先绑定别名避免弃用告警与功能缺失
    sys.modules.setdefault("fitz", pymupdf)
    try:
        from pdf2docx import Converter
    except ImportError:
        if progress_cb:
            progress_cb(-1, "错误: pdf2docx 未安装，无法保留排版转换")
        return False
    if progress_cb:
        progress_cb(30, "解析 PDF…")
    cv = None
    try:
        cv = Converter(input_path)
        if progress_cb:
            progress_cb(60, "生成 Word…")
        cv.convert(output_path)
    except InterruptedError:
        raise
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: 转换失败 - {e}")
        return False
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:  # noqa: BLE001 - 关闭失败不掩盖转换结果
                pass
    if progress_cb:
        progress_cb(100, "生成完成")
    return True


def _ocr_scan_pdf_to_docx(input_path, output_path, lang, dpi,
                          keep_images, table_recognition, progress_cb):
    """扫描/无文本层 PDF：逐页嵌入原图（保真）+ 重建真·Word表格 + OCR文字。

    纯文本 OCR 会丢失表格与图片，对扫描件改为：
    - 嵌入原图（保留图表/版式，所见即所得）；
    - 用 RapidOCR 带坐标重建「真·Word 表格」（单元格可编辑）；
    - 非表格区域的 OCR 文字还原为段落。
    table_recognition=False 时退化为「原图 + 平铺文字」。
    """
    try:
        import pymupdf
    except ImportError:
        if progress_cb:
            progress_cb(-1, "错误: PyMuPDF (fitz) 未安装，无法处理 PDF")
        return False
    try:
        from docx import Document
    except ImportError:
        if progress_cb:
            progress_cb(-1, "错误: python-docx 未安装，无法生成 Word")
        return False

    try:
        pdf = pymupdf.open(input_path)
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: PDF 打开失败 - {e}")
        return False

    total = max(pdf.page_count, 1)
    doc = Document()
    tmpdir = tempfile.mkdtemp(prefix="fm_ocr_")
    try:
        from core.ocr_tool import ocr_image_detailed
        for idx in range(pdf.page_count):
            if progress_cb:
                progress_cb(int((idx + 1) / total * 95),
                            f"识别 PDF 第 {idx + 1}/{total} 页…")
            page = pdf.load_page(idx)
            try:
                pix = page.get_pixmap(dpi=dpi)
                img_path = os.path.join(tmpdir, f"p{idx}.png")
                pix.save(img_path)
            except Exception:  # noqa: BLE001 - 单页渲染失败跳过
                continue

            W, H = pix.width, pix.height
            if table_recognition:
                from core.ocr_tool import ocr_image_raw
                raw = ocr_image_raw(img_path, lang)
                details = (_raw_to_details(raw, W, H) if raw
                           else ocr_image_detailed(img_path, lang))
                boxes = _denormalize_boxes(details, W, H)
                from core.ocr_table import reconstruct_tables_hybrid
                tables, free = reconstruct_tables_hybrid(
                    boxes, W, H, image_path=img_path, dpi=dpi, ocr_raw=raw)
                _build_docx_from_ocr(
                    doc, tables, free,
                    image_to_embed=(img_path if keep_images else None),
                    img_w=W, img_h=H, dpi=dpi, image_path=img_path)
            else:
                if keep_images:
                    from docx.shared import Inches
                    doc.add_picture(img_path, width=Inches(6.3))
                details = ocr_image_detailed(img_path, lang)
                details.sort(key=lambda lb: _sort_key(lb[1]))
                for text, box in details:
                    if text and text.strip():
                        para = doc.add_paragraph()
                        run = para.add_run(text)
                        pt = _box_point_size(_denormalize_box(box, W, H), dpi)
                        if pt > 0:
                            from docx.shared import Pt
                            run.font.size = Pt(pt)

            if idx < pdf.page_count - 1:
                doc.add_page_break()
    finally:
        try:
            pdf.close()
        except Exception:  # noqa: BLE001
            pass
        _cleanup_tmp(tmpdir)

    if progress_cb:
        progress_cb(100, "生成完成")
    try:
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
    except OSError as e:
        if progress_cb:
            progress_cb(-1, f"错误: 无法保存文件 - {e}")
        return False
    return True


def ocr_pdf_to_docx(input_path, output_path, lang="chi_sim+eng",
                    dpi=300, progress_cb=None, keep_images=True,
                    table_recognition=True):
    """高级 OCR：PDF 全文转为可编辑 Word。

    自动判别：
    - 数字 PDF（多数页有文本层）→ pdf2docx 保留表格/排版；
    - 扫描/图片型 PDF（无文本层）→ 逐页嵌入原图 + 重建真·Word表格 + OCR文字。
    返回是否成功。
    """
    try:
        import pymupdf
    except ImportError:
        if progress_cb:
            progress_cb(-1, "错误: PyMuPDF (fitz) 未安装，无法处理 PDF")
        return False

    try:
        pdf = pymupdf.open(input_path)
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: PDF 打开失败 - {e}")
        return False
    try:
        total = pdf.page_count
        if total == 0:
            if progress_cb:
                progress_cb(-1, "错误: PDF 页数为 0")
            return False
        text_pages = 0
        for pg in pdf:
            try:
                if pg.get_text("text").strip():
                    text_pages += 1
            except Exception:  # noqa: BLE001
                continue
    finally:
        try:
            pdf.close()
        except Exception:  # noqa: BLE001
            pass

    # 多数页有文本层 → 数字文档，优先保留表格/排版
    if text_pages / total >= 0.5:
        return _pdf_to_docx_via_pdf2docx(input_path, output_path, progress_cb)
    # 否则视为扫描件：嵌入原图 + 重建真·Word表格 + OCR 文字
    return _ocr_scan_pdf_to_docx(input_path, output_path, lang, dpi,
                                 keep_images, table_recognition, progress_cb)


def _ocr_file_to_docx_impl(input_path, output_path, lang="chi_sim+eng",
                           dpi=300, progress_cb=None, keep_images=True,
                           table_recognition=True):
    """OCR 单个文件（图片或 PDF）转为可编辑 docx。返回是否成功。

    图片：RapidOCR 带坐标重建（真·Word表格 + 自由文本）；
    PDF：走 ocr_pdf_to_docx。
    """
    ext = os.path.splitext(input_path)[1].lower()
    try:
        if ext in PDF_EXTS:
            return ocr_pdf_to_docx(input_path, output_path, lang, dpi,
                                   progress_cb, keep_images, table_recognition)
        # 图片：OCR 带坐标重建真·Word表格
        from docx import Document
        from PIL import Image
        from core.ocr_tool import ocr_image_raw, ocr_image_detailed
        try:
            with Image.open(input_path) as im:
                iw, ih = im.size
        except Exception:  # noqa: BLE001
            iw = ih = 1000
        raw = ocr_image_raw(input_path, lang, progress_cb)
        details = _raw_to_details(raw, iw, ih) if raw else []
        if not details:
            details = ocr_image_detailed(input_path, lang, progress_cb)
        if not details:
            if progress_cb:
                progress_cb(-1, "未识别到文字")
            return False
        boxes = _denormalize_boxes(details, iw, ih)
        doc = Document()
        if table_recognition:
            from core.ocr_table import reconstruct_tables_hybrid
            tables, free = reconstruct_tables_hybrid(
                boxes, iw, ih, image_path=input_path, dpi=dpi, ocr_raw=raw)
            _build_docx_from_ocr(
                doc, tables, free,
                image_to_embed=(input_path if keep_images else None),
                img_w=iw, img_h=ih, dpi=dpi)
        else:
            details.sort(key=lambda lb: _sort_key(lb[1]))
            for text, box in details:
                if text and text.strip():
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    pt = _box_point_size(_denormalize_box(box, iw, ih), dpi)
                    if pt > 0:
                        from docx.shared import Pt
                        run.font.size = Pt(pt)
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return True
    except InterruptedError:
        raise
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: {e}")
        return False


def ocr_file_to_docx(input_path, output_path, lang="chi_sim+eng",
                     dpi=300, progress_cb=None, keep_images=True,
                     table_recognition=True):
    """安全生成 DOCX：完整写入临时文件后再替换目标文件。"""
    if not _validate_conversion(input_path, output_path, ".docx", progress_cb):
        return False
    staged_path = None
    try:
        staged_path = _staged_output_path(output_path, ".docx")
        ok = _ocr_file_to_docx_impl(
            input_path, staged_path, lang, dpi, _scaled_progress(progress_cb),
            keep_images, table_recognition)
        if not ok:
            return False
        if not os.path.isfile(staged_path) or os.path.getsize(staged_path) == 0:
            _report_error(progress_cb, "错误: 生成的 Word 文件为空")
            return False
        if progress_cb:
            progress_cb(95, "正在保存结果…")
        os.replace(staged_path, output_path)
        staged_path = None
        return True
    except InterruptedError:
        raise
    except OSError as e:
        _report_error(progress_cb, f"错误: 无法保存文件 - {e}")
        return False
    finally:
        if staged_path:
            try:
                os.remove(staged_path)
            except OSError:
                pass


def _ocr_file_to_txt_impl(input_path, output_path, lang="chi_sim+eng",
                           progress_cb=None):
    """OCR 单个文件（图片或 PDF）并写出同名 txt，返回是否成功。

    PDF 逐页识别后合并为一个 txt；图片直接识别。
    """
    ext = os.path.splitext(input_path)[1].lower()
    try:
        if ext in PDF_EXTS:
            text = _ocr_pdf(input_path, lang, progress_cb)
        else:
            text = ocr_image(input_path, lang, progress_cb)
    except InterruptedError:
        raise
    except Exception as e:  # noqa: BLE001
        if progress_cb:
            progress_cb(-1, f"错误: {e}")
        return False
    if not text or not text.strip():
        if progress_cb:
            progress_cb(-1, "未识别到文字")
        return False
    try:
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
    except OSError as e:
        if progress_cb:
            progress_cb(-1, f"错误: 无法保存文件 - {e}")
        return False
    return True


def ocr_file_to_txt(input_path, output_path, lang="chi_sim+eng",
                    progress_cb=None):
    """安全生成 TXT：失败或取消时保留目标位置已有文件。"""
    if not _validate_conversion(input_path, output_path, ".txt", progress_cb):
        return False
    staged_path = None
    try:
        staged_path = _staged_output_path(output_path, ".txt")
        ok = _ocr_file_to_txt_impl(
            input_path, staged_path, lang, _scaled_progress(progress_cb))
        if not ok:
            return False
        if not os.path.isfile(staged_path) or os.path.getsize(staged_path) == 0:
            _report_error(progress_cb, "错误: 生成的 TXT 文件为空")
            return False
        if progress_cb:
            progress_cb(95, "正在保存结果…")
        os.replace(staged_path, output_path)
        staged_path = None
        return True
    except InterruptedError:
        raise
    except OSError as e:
        _report_error(progress_cb, f"错误: 无法保存文件 - {e}")
        return False
    finally:
        if staged_path:
            try:
                os.remove(staged_path)
            except OSError:
                pass


def make_runner(task):
    """TaskManager 持久化重建用的 runner 工厂（runner_key="ocr_batch"）。

    task.params 支持: lang（识别语言）、export_fmt（txt/docx）、
    keep_images（是否嵌入原图）、table_recognition（是否重建真表格）。
    """
    def runner(t, prog):
        p = t.params or {}
        if p.get("export_fmt") == "docx":
            return ocr_file_to_docx(
                t.file_path, t.output_path,
                lang=p.get("lang", "chi_sim+eng"),
                keep_images=p.get("keep_images", True),
                table_recognition=p.get("table_recognition", True),
                progress_cb=prog)
        return ocr_file_to_txt(t.file_path, t.output_path,
                               lang=p.get("lang", "chi_sim+eng"),
                               progress_cb=prog)
    return runner
