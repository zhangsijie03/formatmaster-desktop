"""引擎库扩展功能真实测试（2026-08-15）。

覆盖：图片增强/特效、表格图表、docx 图片渲染、
加密扩展（RSA/ECC/GCM/签名/证书）回环验证。

运行:venv/Scripts/python -m pytest tests/test_engine_enhance.py -q
"""
import os
import shutil
import subprocess
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

TMP = tempfile.mkdtemp(prefix="fm_eng_")
FF = None


def _ff():
    global FF
    if FF is None:
        from utils.config import get_ffmpeg_path
        FF = get_ffmpeg_path()
    return FF


@pytest.fixture(scope="module", autouse=True)
def _cleanup():
    yield
    shutil.rmtree(TMP, ignore_errors=True)


@pytest.fixture(scope="module")
def sample_png():
    from PIL import Image
    p = os.path.join(TMP, "sample.png")
    Image.new("RGB", (60, 40), (200, 100, 50)).save(p)
    return p


@pytest.fixture(scope="module")
def sample_docx():
    from docx import Document
    from docx.shared import Inches
    from PIL import Image
    img = os.path.join(TMP, "docx_img.png")
    Image.new("RGB", (300, 150), (80, 120, 200)).save(img)
    doc = Document()
    doc.add_heading("带图文档", level=1)
    doc.add_paragraph("正文。")
    doc.add_picture(img, width=Inches(2))
    p = os.path.join(TMP, "sample.docx")
    doc.save(p)
    return p


class TestImageEnhance:
    def test_enhance_and_effects(self, sample_png):
        from core.image_converter import ImageConverter
        cv = ImageConverter()
        for name, kw in [
            ("enhance", dict(contrast=1.3, saturation=1.5, sharpness=2.0)),
            ("hflip", dict(effect="hflip")),
            ("vflip", dict(effect="vflip")),
            ("invert", dict(effect="invert")),
            ("emboss", dict(effect="emboss")),
            ("edges", dict(effect="edges")),
            ("sharpen", dict(effect="sharpen")),
        ]:
            out = os.path.join(TMP, f"fx_{name}.png")
            assert cv.convert(sample_png, out, 90, **kw)
            assert os.path.getsize(out) > 0

    def test_chinese_watermark_on_grayscale_image(self, tmp_path):
        """macOS 中文字体回退与灰度绘制必须真实产生水印。"""
        from PIL import Image, ImageChops
        from core.image_converter import ImageConverter

        source = tmp_path / "gray.png"
        output = tmp_path / "watermarked.png"
        original = Image.new("L", (640, 360), 128)
        original.save(source)

        assert ImageConverter().convert(
            str(source), str(output), watermark_text="中文水印",
            watermark_position="左上角", grayscale=True)
        with Image.open(output) as rendered:
            assert ImageChops.difference(original, rendered.convert("L")).getbbox()

    def test_exif_preserved_unless_privacy_strip_enabled(self, tmp_path):
        """默认保留 EXIF；用户明确开启隐私清理后再删除。"""
        from PIL import Image
        from core.image_converter import ImageConverter

        source = tmp_path / "source.jpg"
        kept = tmp_path / "kept.jpg"
        stripped = tmp_path / "stripped.jpg"
        exif = Image.Exif()
        exif[0x010F] = "FormatMaster Camera"
        Image.new("RGB", (80, 60), "navy").save(source, exif=exif)

        converter = ImageConverter()
        assert converter.convert(str(source), str(kept))
        assert converter.convert(str(source), str(stripped), strip_exif=True)
        with Image.open(kept) as image:
            assert image.getexif().get(0x010F) == "FormatMaster Camera"
        with Image.open(stripped) as image:
            assert image.getexif().get(0x010F) is None





class TestTableChart:
    def test_charts(self):
        import core.table_recognizer as tr
        orig = tr.recognize_rows
        tr.recognize_rows = lambda path, cb=None: [
            ["月份", "销量", "利润"], ["一月", 120, 30],
            ["二月", 150, 40], ["三月", 110, 25]]
        try:
            for ct in (None, "bar", "line", "pie"):
                out = os.path.join(TMP, f"tbl_{ct or 'none'}.xlsx")
                assert tr.table_to_xlsx("dummy", out, None, chart_type=ct)
                assert os.path.getsize(out) > 0
        finally:
            tr.recognize_rows = orig


class TestDocxImageRender:
    def test_image_rendered(self, sample_docx):
        from core.doc_office_pdf import render_docx_to_pdf
        out = os.path.join(TMP, "with_img.pdf")
        assert render_docx_to_pdf(sample_docx, out)
        import pymupdf
        total = sum(len(p.get_images()) for p in pymupdf.open(out))
        assert total >= 1, "docx 图片未渲染进 PDF"


class TestCryptoAdvanced:
    def _src(self):
        p = os.path.join(TMP, "data.bin")
        with open(p, "wb") as f:
            f.write(os.urandom(200000))
        return p

    def test_rsa_roundtrip(self):
        from core.crypto_advanced import (decrypt_asymmetric,
                                          encrypt_asymmetric, generate_keypair)
        src = self._src()
        priv, pub = generate_keypair("rsa")
        enc = os.path.join(TMP, "rsa.bin")
        dec = os.path.join(TMP, "rsa_out.bin")
        assert encrypt_asymmetric(src, enc, pub)
        assert decrypt_asymmetric(enc, dec, priv)
        assert open(src, "rb").read() == open(dec, "rb").read()

    def test_ecc_roundtrip(self):
        from core.crypto_advanced import (decrypt_asymmetric,
                                          encrypt_asymmetric, generate_keypair)
        src = self._src()
        priv, pub = generate_keypair("ecc")
        enc = os.path.join(TMP, "ecc.bin")
        dec = os.path.join(TMP, "ecc_out.bin")
        assert encrypt_asymmetric(src, enc, pub)
        assert decrypt_asymmetric(enc, dec, priv)
        assert open(src, "rb").read() == open(dec, "rb").read()

    def test_sign_verify(self):
        from core.crypto_advanced import (generate_keypair, sign_file,
                                          verify_signature)
        src = self._src()
        priv, pub = generate_keypair("rsa")
        sig = sign_file(src, priv)
        assert sig
        assert verify_signature(src, pub, sig)[0]
        with open(src, "ab") as f:
            f.write(b"tamper")
        assert not verify_signature(src, pub, sig)[0]

    def test_gcm_roundtrip(self):
        from core.crypto_advanced import (decrypt_file_gcm,
                                          encrypt_file_gcm)
        src = self._src()
        enc = os.path.join(TMP, "gcm.bin")
        dec = os.path.join(TMP, "gcm_out.bin")
        assert encrypt_file_gcm(src, enc, "密码123")
        assert decrypt_file_gcm(enc, dec, "密码123")
        assert open(src, "rb").read() == open(dec, "rb").read()

    def test_cert(self):
        from core.crypto_advanced import generate_self_signed_cert
        cert, key = generate_self_signed_cert(
            "fm.test", os.path.join(TMP, "cert.crt"), days=30)
        assert cert and os.path.isfile(cert) and os.path.isfile(key)
