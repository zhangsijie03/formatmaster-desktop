"""高级 OCR 单元测试：PDF/图片 → 可编辑 docx 混合管线。

不依赖真实 RapidOCR 模型（离线），扫描路径用 monkeypatch 替换
core.ocr_tool.ocr_image_detailed 进行验证。
"""
import os

import pytest


def _make_text_pdf(path, pages=1):
    """构造含文本层的 PDF（ASCII，避免字体依赖）。"""
    import pymupdf
    doc = pymupdf.open()
    for i in range(pages):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page{i + 1} Hello World")
    doc.save(path)
    doc.close()


def _make_blank_pdf(path, pages=1):
    """构造无文本层的 PDF（空白页 → 触发扫描件路径）。"""
    import pymupdf
    doc = pymupdf.open()
    for _ in range(pages):
        doc.new_page()
    doc.save(path)
    doc.close()


def _docx_text(path):
    from docx import Document
    d = Document(path)
    return "\n".join(p.text for p in d.paragraphs)


def test_ocr_pdf_to_docx_digital_dispatch(monkeypatch, tmp_path):
    """数字 PDF：走 pdf2docx 保留表格，不触发 OCR。"""
    from core import ocr_batch
    pdf = tmp_path / "in.pdf"
    _make_text_pdf(str(pdf))
    out = tmp_path / "out.docx"

    calls = {"pdf2docx": 0, "ocr": False}

    def fake_pdf2docx(inp, op, progress_cb=None):
        calls["pdf2docx"] += 1
        from docx import Document
        Document().save(op)  # 占位 docx
        return True

    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        calls["ocr"] = True
        return []

    monkeypatch.setattr(ocr_batch, "_pdf_to_docx_via_pdf2docx", fake_pdf2docx)
    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_pdf_to_docx(str(pdf), str(out), progress_cb=None)
    assert ok is True
    assert calls["pdf2docx"] == 1
    assert calls["ocr"] is False  # 数字文档不应走 OCR
    assert out.exists()


def test_ocr_pdf_to_docx_scan_embed(monkeypatch, tmp_path):
    """扫描件：逐页嵌入原图 + OCR 文字，docx 含图片且文字按 y 排序。"""
    from core import ocr_batch
    pdf = tmp_path / "scan.pdf"
    _make_blank_pdf(str(pdf), pages=2)  # 2 页无文本 → 扫描件
    out = tmp_path / "out.docx"

    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        return [("B", (0.0, 0.2, 1.0, 0.3)),
                ("A", (0.0, 0.0, 1.0, 0.1))]

    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_pdf_to_docx(str(pdf), str(out), dpi=100,
                                   progress_cb=None)
    assert ok is True
    from docx import Document
    d = Document(str(out))
    assert len(d.inline_shapes) == 2  # 每页嵌入 1 张原图
    text = _docx_text(str(out))
    assert "A" in text and "B" in text
    # 按版面 y 坐标排序：A(y=0.0) 应在 B(y=0.2) 之前
    assert text.index("A") < text.index("B")


def test_ocr_file_to_docx_image(monkeypatch, tmp_path):
    """图片：直接 OCR 带坐标重建 docx。"""
    from core import ocr_batch
    from PIL import Image
    img = tmp_path / "t.png"
    Image.new("RGB", (60, 60), "white").save(str(img))
    out = tmp_path / "out.docx"

    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        return [("ImgLine1", (0.0, 0.0, 1.0, 0.1)),
                ("ImgLine2", (0.0, 0.2, 1.0, 0.3))]

    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_file_to_docx(str(img), str(out), progress_cb=None)
    assert ok is True
    text = _docx_text(str(out))
    assert "ImgLine1" in text and "ImgLine2" in text


def test_ocr_file_to_docx_missing_input():
    """输入文件不存在应返回 False。"""
    from core import ocr_batch
    ok = ocr_batch.ocr_file_to_docx("__no_such_file__.pdf", "out.docx")
    assert ok is False


def test_make_runner_dispatch(monkeypatch):
    """make_runner 按 export_fmt 分发到 txt / docx 实现。"""
    from core import ocr_batch

    class Task:
        def __init__(self, params):
            self.params = params
            self.file_path = "f.pdf"
            self.output_path = "o.docx"

    txt_log, docx_log = [], []

    def fake_txt(fp, op, lang="chi_sim+eng", progress_cb=None, **kw):
        txt_log.append(1)
        return True

    def fake_docx(fp, op, lang="chi_sim+eng", dpi=300, progress_cb=None, **kw):
        docx_log.append(1)
        return True

    monkeypatch.setattr(ocr_batch, "ocr_file_to_txt", fake_txt)
    monkeypatch.setattr(ocr_batch, "ocr_file_to_docx", fake_docx)

    r_txt = ocr_batch.make_runner(Task({"lang": "chi_sim+eng", "export_fmt": "txt"}))
    r_txt(Task({"lang": "chi_sim+eng", "export_fmt": "txt"}), None)
    assert txt_log and not docx_log

    r_docx = ocr_batch.make_runner(Task({"lang": "chi_sim+eng", "export_fmt": "docx"}))
    r_docx(Task({"lang": "chi_sim+eng", "export_fmt": "docx"}), None)
    assert docx_log


def test_ocr_image_detailed_shape(monkeypatch, tmp_path):
    """ocr_image_detailed 返回 [(text, box)]，box 为归一化 0~1。"""
    from core import ocr_tool
    from PIL import Image
    img = tmp_path / "t.png"
    Image.new("RGB", (100, 50), "white").save(str(img))

    class FakeEngine:
        def __call__(self, p):
            return [[[[10, 5], [40, 5], [40, 20], [10, 20]], "Hi", 0.99]], 0.01

    monkeypatch.setattr(ocr_tool, "_get_engine", lambda: FakeEngine())

    res = ocr_tool.ocr_image_detailed(str(img))
    assert len(res) == 1
    text, box = res[0]
    assert text == "Hi"
    assert all(0.0 <= v <= 1.0 for v in box)
    assert 0.05 < box[1] < 0.5


def test_ocr_file_to_docx_image_table(monkeypatch, tmp_path):
    """图片 + 表格识别：生成真·Word表格（单元格可编辑），内容落位正确。"""
    from core import ocr_batch
    from PIL import Image
    from docx import Document
    img = tmp_path / "grid.png"
    Image.new("RGB", (400, 400), "white").save(str(img))
    out = tmp_path / "out.docx"

    # 归一化 2x2 表格（与 400x400 图像同尺寸）
    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        return [("Name", (0.10, 0.10, 0.30, 0.15)),
                ("Age", (0.40, 0.10, 0.60, 0.15)),
                ("Bob", (0.10, 0.30, 0.30, 0.35)),
                ("30", (0.40, 0.30, 0.60, 0.35))]

    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_file_to_docx(
        str(img), str(out), keep_images=False,
        table_recognition=True, progress_cb=None)
    assert ok is True
    d = Document(str(out))
    assert len(d.tables) == 1
    tbl = d.tables[0]
    assert len(tbl.rows) == 2 and len(tbl.columns) == 2
    assert tbl.cell(0, 0).text == "Name"
    assert tbl.cell(0, 1).text == "Age"
    assert tbl.cell(1, 0).text == "Bob"
    assert tbl.cell(1, 1).text == "30"


def test_ocr_file_to_docx_image_no_table(monkeypatch, tmp_path):
    """表格识别关闭：图片只输出平铺文字，不生成表格。"""
    from core import ocr_batch
    from PIL import Image
    from docx import Document
    img = tmp_path / "plain.png"
    Image.new("RGB", (200, 200), "white").save(str(img))
    out = tmp_path / "out.docx"

    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        return [("Plain1", (0.0, 0.0, 1.0, 0.1)),
                ("Plain2", (0.0, 0.2, 1.0, 0.3))]

    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_file_to_docx(
        str(img), str(out), keep_images=False,
        table_recognition=False, progress_cb=None)
    assert ok is True
    d = Document(str(out))
    assert len(d.tables) == 0
    assert "Plain1" in "\n".join(p.text for p in d.paragraphs)


def test_extract_docx_text_includes_tables(tmp_path):
    """extract_docx_text 应同时提取段落与表格单元格文字。"""
    from core import ocr_batch
    from docx import Document
    out = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("段落文字")
    tbl = d.add_table(1, 2)
    tbl.cell(0, 0).text = "单元格A"
    tbl.cell(0, 1).text = "单元格B"
    d.save(str(out))
    text = ocr_batch.extract_docx_text(str(out))
    assert "段落文字" in text
    assert "单元格A" in text and "单元格B" in text


def test_ocr_docx_applies_font_sizes(monkeypatch, tmp_path):
    """高级 OCR 生成的 docx 应把字号设为接近原文件的磅值（非默认 11pt）。

    用 12pt 设计高度的假框（400px 图、dpi=300）：20px→4pt? 这里用
    0.1176 高度 ≈ 12pt，断言单元格 run 字号被显式设置且落在合理区间。
    """
    from core import ocr_batch
    from PIL import Image
    from docx import Document
    img = tmp_path / "font.png"
    Image.new("RGB", (400, 400), "white").save(str(img))
    out = tmp_path / "out.docx"

    def fake_detailed(p, lang="chi_sim+eng", progress_cb=None):
        # 行高 0.1176 * 400 = 47px @dpi300*calib0.85 ≈ 12pt
        h = 0.1176
        return [("Name", (0.10, 0.10, 0.30, 0.10 + h)),
                ("Age", (0.40, 0.10, 0.60, 0.10 + h)),
                ("Bob", (0.10, 0.30, 0.30, 0.30 + h)),
                ("30", (0.40, 0.30, 0.60, 0.30 + h))]

    monkeypatch.setattr("core.ocr_tool.ocr_image_detailed", fake_detailed)

    ok = ocr_batch.ocr_file_to_docx(
        str(img), str(out), keep_images=False,
        table_recognition=True, progress_cb=None)
    assert ok is True
    d = Document(str(out))
    cell_sizes = []
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.size is not None:
                            cell_sizes.append(r.font.size.pt)
    assert cell_sizes, "单元格未设置字号"
    # 设计值 ≈12pt，允许校准误差
    assert all(8.0 <= s <= 16.0 for s in cell_sizes)


def test_ocr_docx_merges_cells(tmp_path):
    """_add_word_table 应把 table['merges'] 写成 Word 合并单元格。"""
    from core import ocr_batch
    from docx import Document
    from docx.oxml.ns import qn

    table = {
        "bbox": (50, 100, 420, 340),
        "rows": 3, "cols": 2,
        "grid": [["总标题", ""], ["a1", "b1"], ["a2", "b2"]],
        "sizes": [[10.0, 0.0], [9.0, 9.0], [9.0, 9.0]],
        "merges": [(0, 0, 0, 1)],
        "col_bounds": [(50, 235), (235, 420)],
    }
    doc = Document()
    ocr_batch._add_word_table(doc, table)
    tbl = doc.tables[0]
    # 至少存在一个合并标记 w:gridSpan
    gridspans = list(tbl._tbl.iter(qn("w:gridSpan")))
    assert gridspans, "未生成合并单元格"
    assert gridspans[0].get(qn("w:val")) == "2"


def test_ocr_docx_col_widths_set(tmp_path):
    """_add_word_table 应固定表格宽度并按 col_bounds 比例设列宽。"""
    from core import ocr_batch
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = {
        "bbox": (50, 100, 420, 340),
        "rows": 2, "cols": 2,
        "grid": [["a", "b"], ["c", "d"]],
        "sizes": [[10.0, 10.0], [10.0, 10.0]],
        "merges": [],
        "col_bounds": [(50, 150), (150, 420)],  # 左窄右宽
    }
    doc = Document()
    ocr_batch._add_word_table(doc, table)
    tbl = doc.tables[0]
    tblPr = tbl._tbl.tblPr
    w = tblPr.find(qn("w:tblW"))
    assert w is not None and w.get(qn("w:type")) == "dxa"
    # 两列宽度不同（左窄右宽）
    assert tbl.columns[0].width is not None
    assert tbl.columns[1].width is not None
    w0 = tbl.columns[0].width.pt if hasattr(tbl.columns[0].width, "pt") else 0
    w1 = tbl.columns[1].width.pt if hasattr(tbl.columns[1].width, "pt") else 0
    assert w1 > w0  # 右列更宽

